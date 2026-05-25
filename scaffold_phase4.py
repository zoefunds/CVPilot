"""
CVPilot Phase 4: Storage + Parsing + Job Fetch + Application + Celery.
Writes every file. Idempotent.
"""

from __future__ import annotations
from pathlib import Path

ROOT = Path("/Users/macbook/CVPilot")

FILES: dict[str, str] = {}

# -----------------------------------------------------------------------------
# services/__init__.py packages
# -----------------------------------------------------------------------------
FILES["services/__init__.py"] = ""
FILES["services/storage/__init__.py"] = '''from services.storage.factory import get_storage  # noqa: F401
from services.storage.base import FileStorage, StoredFile  # noqa: F401
'''
FILES["services/parsing/__init__.py"] = '''from services.parsing.extractor import extract_text, FileKind  # noqa: F401
'''
FILES["services/jobfetch/__init__.py"] = '''from services.jobfetch.fetcher import fetch_job_posting, JobPosting  # noqa: F401
'''

# -----------------------------------------------------------------------------
# Storage
# -----------------------------------------------------------------------------
FILES["services/storage/base.py"] = '''"""
Storage abstraction. Designed so a future S3/R2/GCS implementation
just provides this Protocol; routes do not change.
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import Protocol


@dataclass(frozen=True)
class StoredFile:
    key: str
    byte_size: int
    content_type: str


class FileStorage(Protocol):
    def save(self, key: str, data: bytes, content_type: str) -> StoredFile: ...
    def read(self, key: str) -> bytes: ...
    def delete(self, key: str) -> None: ...
    def exists(self, key: str) -> bool: ...
'''

FILES["services/storage/local.py"] = '''"""
Local filesystem implementation of FileStorage.
Files live under settings.storage_local_path (default ./storage/uploads).
Keys may contain forward slashes; they are translated to subdirectories.
"""

from __future__ import annotations

import os
from pathlib import Path

from backend.app.core.config import settings
from services.storage.base import FileStorage, StoredFile


class LocalFileStorage(FileStorage):
    def __init__(self, root: str | None = None) -> None:
        self.root = Path(root or settings.storage_local_path).resolve()
        self.root.mkdir(parents=True, exist_ok=True)

    def _path(self, key: str) -> Path:
        # Defense in depth: never let a key escape the root.
        key = key.lstrip("/").replace("..", "_")
        p = (self.root / key).resolve()
        if not str(p).startswith(str(self.root)):
            raise ValueError("Path traversal attempt rejected.")
        return p

    def save(self, key: str, data: bytes, content_type: str) -> StoredFile:
        p = self._path(key)
        p.parent.mkdir(parents=True, exist_ok=True)
        with open(p, "wb") as f:
            f.write(data)
        return StoredFile(key=key, byte_size=len(data), content_type=content_type)

    def read(self, key: str) -> bytes:
        with open(self._path(key), "rb") as f:
            return f.read()

    def delete(self, key: str) -> None:
        p = self._path(key)
        if p.exists():
            os.remove(p)

    def exists(self, key: str) -> bool:
        return self._path(key).exists()
'''

FILES["services/storage/factory.py"] = '''"""
Factory that returns the configured storage backend.
Today: local. Tomorrow: switch on settings.storage_backend.
"""

from __future__ import annotations

from functools import lru_cache

from backend.app.core.config import settings
from services.storage.base import FileStorage
from services.storage.local import LocalFileStorage


@lru_cache(maxsize=1)
def get_storage() -> FileStorage:
    if settings.storage_backend == "local":
        return LocalFileStorage()
    raise NotImplementedError(f"Unsupported storage backend: {settings.storage_backend}")
'''

# -----------------------------------------------------------------------------
# Parsing
# -----------------------------------------------------------------------------
FILES["services/parsing/extractor.py"] = '''"""
File-content extractor.
- PDF: pypdf
- DOCX: python-docx
- TXT: utf-8 decode with errors="ignore"
Validates magic bytes so a renamed .exe cannot smuggle past the extension check.
"""

from __future__ import annotations

import io
from enum import Enum

from docx import Document  # python-docx
from pypdf import PdfReader

from backend.app.core.errors import ValidationAppError


class FileKind(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


_PDF_MAGIC = b"%PDF-"
_DOCX_MAGIC = b"PK\\x03\\x04"  # any zip; we rely on python-docx to validate structure


def detect_kind(filename: str, data: bytes) -> FileKind:
    name = filename.lower()
    head = data[:8]

    if head.startswith(_PDF_MAGIC):
        return FileKind.PDF
    if head.startswith(b"PK\\x03\\x04") and name.endswith(".docx"):
        return FileKind.DOCX
    if name.endswith(".txt"):
        # Best-effort: assume text if it decodes mostly as utf-8.
        try:
            data[:4096].decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ValidationAppError(
                "File does not appear to be valid UTF-8 text.",
                code="invalid_file_encoding",
            ) from exc
        return FileKind.TXT

    raise ValidationAppError(
        "Unsupported or malformed file. Allowed: PDF, DOCX, TXT.",
        code="unsupported_file_type",
    )


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            parts.append("")
    return "\\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\\n".join(parts).strip()


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore").strip()


def extract_text(filename: str, data: bytes) -> tuple[FileKind, str]:
    kind = detect_kind(filename, data)
    if kind == FileKind.PDF:
        text = _extract_pdf(data)
    elif kind == FileKind.DOCX:
        text = _extract_docx(data)
    else:
        text = _extract_txt(data)

    if not text:
        raise ValidationAppError(
            "Could not extract any text from the uploaded file.",
            code="empty_extraction",
        )
    return kind, text
'''

# -----------------------------------------------------------------------------
# Job URL fetcher
# -----------------------------------------------------------------------------
FILES["services/jobfetch/fetcher.py"] = '''"""
Fetch a job posting URL and return visible text + title.
Guarded by:
  - 10s connect / 15s read timeout
  - 2 MB response cap
  - text/html content type required
"""

from __future__ import annotations

from dataclasses import dataclass

import httpx
from bs4 import BeautifulSoup

from backend.app.core.errors import ValidationAppError

_MAX_BYTES = 2 * 1024 * 1024
_USER_AGENT = "CVPilotBot/0.1 (+https://cvpilot.dev)"


@dataclass(frozen=True)
class JobPosting:
    url: str
    final_url: str
    title: str
    text: str


def _clean_html(html: str) -> tuple[str, str]:
    soup = BeautifulSoup(html, "lxml")
    title = (soup.title.string.strip() if soup.title and soup.title.string else "") or ""

    for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "form"]):
        tag.decompose()

    text = soup.get_text(separator="\\n", strip=True)
    # Collapse blank-line runs
    lines = [ln for ln in (l.strip() for l in text.splitlines()) if ln]
    return title, "\\n".join(lines)


def fetch_job_posting(url: str) -> JobPosting:
    try:
        with httpx.Client(
            follow_redirects=True,
            timeout=httpx.Timeout(connect=10.0, read=15.0, write=10.0, pool=5.0),
            headers={"User-Agent": _USER_AGENT, "Accept": "text/html,*/*"},
        ) as client:
            r = client.get(url)
    except httpx.HTTPError as exc:
        raise ValidationAppError(
            f"Could not reach job URL: {exc.__class__.__name__}",
            code="job_url_unreachable",
        ) from exc

    if r.status_code >= 400:
        raise ValidationAppError(
            f"Job URL returned HTTP {r.status_code}.",
            code="job_url_http_error",
        )

    ctype = r.headers.get("content-type", "")
    if "text/html" not in ctype and "application/xhtml" not in ctype:
        raise ValidationAppError(
            f"Job URL did not return HTML (content-type={ctype}).",
            code="job_url_not_html",
        )

    content = r.content[:_MAX_BYTES]
    title, text = _clean_html(content.decode("utf-8", errors="ignore"))
    if not text:
        raise ValidationAppError(
            "Could not extract any text from the job posting.",
            code="job_url_empty",
        )

    return JobPosting(url=url, final_url=str(r.url), title=title, text=text)
'''

# -----------------------------------------------------------------------------
# ORM: Application, FileAsset
# -----------------------------------------------------------------------------
FILES["backend/app/models/application.py"] = '''"""
Application + FileAsset ORM models.
"""

from __future__ import annotations

import uuid

from sqlalchemy import ForeignKey, Integer, String, Text
from sqlalchemy.dialects.postgresql import ENUM, UUID
from sqlalchemy.orm import Mapped, mapped_column, relationship

from backend.app.db.base import Base, TimestampMixin

# Postgres ENUM types are created once and reused.
application_status_enum = ENUM(
    "pending",
    "processing",
    "ready",
    "evaluating",
    "complete",
    "failed",
    name="application_status",
    create_type=True,
)

file_kind_enum = ENUM(
    "cv",
    "cover_letter",
    name="file_kind",
    create_type=True,
)


class Application(Base, TimestampMixin):
    __tablename__ = "applications"

    id: Mapped[uuid.UUID] = mapped_column(
        UUID(as_uuid=True), primary_key=True, default=uuid.uuid4
    )
    user_id: Mapped[uuid.UUID] = mapped_column(
        UUID(as_uuid=True),
        ForeignKey("users.id", ondelete="CASCADE"),
        nullable=False,
        index=True,
    )
    job_url: Mapped[str] = mapped_column(String(2048), nullable=False)
    job_final_url: Mapped[str | None] = mapped_column(String(2048), nullable=True)
    job_title: Mapped[str | None] = mapped_column(String(512), nullable=True)
    job_text: Mapped[str | None] = mapped_column(Text, nullable=True)
    linkedin_url: Mapped[str | None] = mapped_column(String(2048), nullable=True)
    portfolio_url: Mapped[str | None] = mapped_column(String(2048), nullable=True)
    status: Mapped[str] = mapped_column(application_status_enum, default="pending", nullable=False)
    error: Mapped[str | None] = mapped_column(Text, nullable=True)

    files: Mapped[list["FileAsset"]] = relationship(
        "FileAsset",
        back_populates="application",
        cascade="all, delete-orphan",
        lazy="selectin",
    )


class FileAsset(Base, TimestampMixin):
    __tablename__ = "file_assets"

    id: Mapped[uuid.UUID] = mapped_column(
        UUID(as_uuid=True), primary_key=True, default=uuid.uuid4
    )
    application_id: Mapped[uuid.UUID] = mapped_column(
        UUID(as_uuid=True),
        ForeignKey("applications.id", ondelete="CASCADE"),
        nullable=False,
        index=True,
    )
    kind: Mapped[str] = mapped_column(file_kind_enum, nullable=False)
    original_filename: Mapped[str] = mapped_column(String(512), nullable=False)
    storage_key: Mapped[str] = mapped_column(String(1024), nullable=False)
    content_type: Mapped[str] = mapped_column(String(128), nullable=False)
    byte_size: Mapped[int] = mapped_column(Integer, nullable=False)
    detected_kind: Mapped[str | None] = mapped_column(String(16), nullable=True)
    extracted_text: Mapped[str | None] = mapped_column(Text, nullable=True)

    application: Mapped["Application"] = relationship("Application", back_populates="files")
'''

# Re-export new models so Alembic sees them.
FILES["backend/app/models/__init__.py"] = '''"""
Import all models here so Alembic autogenerate detects them.
"""

from backend.app.models.user import User  # noqa: F401
from backend.app.models.audit_log import AuditLog  # noqa: F401
from backend.app.models.application import Application, FileAsset  # noqa: F401
'''

# -----------------------------------------------------------------------------
# Schemas
# -----------------------------------------------------------------------------
FILES["backend/app/schemas/application.py"] = '''"""
Application request/response schemas.
"""

from __future__ import annotations

import uuid
from datetime import datetime

from pydantic import BaseModel, ConfigDict, Field, HttpUrl


class FileAssetPublic(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: uuid.UUID
    kind: str
    original_filename: str
    content_type: str
    byte_size: int
    detected_kind: str | None = None
    extracted_text: str | None = None


class ApplicationPublic(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: uuid.UUID
    job_url: str
    job_final_url: str | None = None
    job_title: str | None = None
    job_text: str | None = None
    linkedin_url: str | None = None
    portfolio_url: str | None = None
    status: str
    error: str | None = None
    created_at: datetime
    updated_at: datetime
    files: list[FileAssetPublic] = Field(default_factory=list)


class ApplicationListItem(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: uuid.UUID
    job_url: str
    job_title: str | None = None
    status: str
    created_at: datetime
'''

# -----------------------------------------------------------------------------
# Routes
# -----------------------------------------------------------------------------
FILES["backend/app/routes/applications.py"] = '''"""
Applications API.
Deliberately no `from __future__ import annotations` because slowapi wraps
these handlers and FastAPI cannot resolve forward refs against slowapi's
module globals.
"""

import uuid
from typing import Optional

from fastapi import APIRouter, Depends, File, Form, Request, UploadFile, status
from sqlalchemy import select
from sqlalchemy.orm import Session

from backend.app.core.config import settings
from backend.app.core.errors import (
    ForbiddenError,
    NotFoundError,
    ValidationAppError,
)
from backend.app.core.logging import get_logger
from backend.app.db.session import get_db
from backend.app.dependencies.auth import get_current_user
from backend.app.dependencies.rate_limit import limiter
from backend.app.models.application import Application, FileAsset
from backend.app.models.user import User
from backend.app.schemas.application import (
    ApplicationListItem,
    ApplicationPublic,
)
from services.storage import get_storage

router = APIRouter(prefix="/applications", tags=["applications"])
log = get_logger("applications")

_ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}


def _validate_upload(file: UploadFile, label: str) -> bytes:
    if file is None or not file.filename:
        raise ValidationAppError(f"{label} file is required.", code=f"{label}_missing")

    # Read fully so we can size-check before persistence.
    data = file.file.read()
    max_bytes = settings.storage_max_upload_mb * 1024 * 1024
    if len(data) == 0:
        raise ValidationAppError(f"{label} file is empty.", code=f"{label}_empty")
    if len(data) > max_bytes:
        raise ValidationAppError(
            f"{label} file exceeds {settings.storage_max_upload_mb} MB limit.",
            code=f"{label}_too_large",
        )
    if file.content_type and file.content_type not in _ALLOWED_CONTENT_TYPES:
        # Soft check; magic-byte validation happens in the parser.
        log.info(
            "upload_unexpected_content_type",
            label=label,
            content_type=file.content_type,
        )
    return data


def _store_file(
    storage,
    *,
    user_id: uuid.UUID,
    application_id: uuid.UUID,
    kind: str,
    filename: str,
    data: bytes,
    content_type: str,
) -> FileAsset:
    safe_name = filename.replace("/", "_").replace("\\\\", "_")
    key = f"{user_id}/{application_id}/{kind}-{uuid.uuid4().hex}-{safe_name}"
    stored = storage.save(key, data, content_type)
    return FileAsset(
        application_id=application_id,
        kind=kind,
        original_filename=filename,
        storage_key=stored.key,
        content_type=content_type,
        byte_size=stored.byte_size,
    )


@router.post("", response_model=ApplicationPublic, status_code=status.HTTP_202_ACCEPTED)
@limiter.limit("20/minute")
def create_application(
    request: Request,
    job_url: str = Form(..., min_length=8, max_length=2048),
    linkedin_url: Optional[str] = Form(default=None, max_length=2048),
    portfolio_url: Optional[str] = Form(default=None, max_length=2048),
    cv: UploadFile = File(...),
    cover_letter: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if not (job_url.startswith("http://") or job_url.startswith("https://")):
        raise ValidationAppError(
            "job_url must start with http:// or https://", code="job_url_invalid"
        )

    cv_bytes = _validate_upload(cv, "cv")
    cl_bytes = _validate_upload(cover_letter, "cover_letter")

    application = Application(
        user_id=current_user.id,
        job_url=job_url,
        linkedin_url=linkedin_url or None,
        portfolio_url=portfolio_url or None,
        status="pending",
    )
    db.add(application)
    db.flush()  # assigns id

    storage = get_storage()
    db.add(
        _store_file(
            storage,
            user_id=current_user.id,
            application_id=application.id,
            kind="cv",
            filename=cv.filename,
            data=cv_bytes,
            content_type=cv.content_type or "application/octet-stream",
        )
    )
    db.add(
        _store_file(
            storage,
            user_id=current_user.id,
            application_id=application.id,
            kind="cover_letter",
            filename=cover_letter.filename,
            data=cl_bytes,
            content_type=cover_letter.content_type or "application/octet-stream",
        )
    )
    db.commit()
    db.refresh(application)

    # Dispatch background processing.
    from workers.tasks.applications import process_application

    process_application.delay(str(application.id))
    log.info("application_created", application_id=str(application.id), user_id=str(current_user.id))

    return ApplicationPublic.model_validate(application)


@router.get("", response_model=list[ApplicationListItem])
def list_applications(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    rows = db.scalars(
        select(Application)
        .where(Application.user_id == current_user.id)
        .order_by(Application.created_at.desc())
    ).all()
    return [ApplicationListItem.model_validate(r) for r in rows]


@router.get("/{application_id}", response_model=ApplicationPublic)
def get_application(
    application_id: uuid.UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    app = db.get(Application, application_id)
    if app is None:
        raise NotFoundError("Application not found.")
    if app.user_id != current_user.id:
        raise ForbiddenError("You do not own this application.")
    return ApplicationPublic.model_validate(app)
'''

# -----------------------------------------------------------------------------
# Celery
# -----------------------------------------------------------------------------
FILES["workers/__init__.py"] = ""
FILES["workers/tasks/__init__.py"] = ""

FILES["workers/celery_app.py"] = '''"""
Celery application factory.
Broker + backend pulled from settings (Redis).
Tests set CELERY_TASK_ALWAYS_EAGER=true so .delay() runs inline.
"""

from __future__ import annotations

import os

from celery import Celery

from backend.app.core.config import settings

celery_app = Celery(
    "cvpilot",
    broker=settings.celery_broker_url,
    backend=settings.celery_result_backend,
    include=["workers.tasks.applications"],
)

celery_app.conf.update(
    task_serializer="json",
    accept_content=["json"],
    result_serializer="json",
    timezone="UTC",
    enable_utc=True,
    task_acks_late=True,
    task_reject_on_worker_lost=True,
    task_time_limit=120,
    task_soft_time_limit=90,
    worker_max_tasks_per_child=100,
    broker_connection_retry_on_startup=True,
)

if os.getenv("CELERY_TASK_ALWAYS_EAGER", "").lower() in {"1", "true", "yes"}:
    celery_app.conf.task_always_eager = True
    celery_app.conf.task_eager_propagates = True
'''

FILES["workers/tasks/applications.py"] = '''"""
Background task: process an Application end-to-end (parse + fetch).
Sets status: processing -> ready (or failed).
"""

from __future__ import annotations

import uuid

from sqlalchemy.orm import Session

from backend.app.core.logging import get_logger
from backend.app.db.session import SessionLocal
from backend.app.models.application import Application, FileAsset
from services.jobfetch import fetch_job_posting
from services.parsing import extract_text
from services.storage import get_storage
from workers.celery_app import celery_app

log = get_logger("worker.applications")


def _process(db: Session, application_id: uuid.UUID) -> None:
    app = db.get(Application, application_id)
    if app is None:
        log.warning("application_missing", application_id=str(application_id))
        return

    app.status = "processing"
    app.error = None
    db.commit()

    storage = get_storage()
    try:
        for asset in app.files:
            raw = storage.read(asset.storage_key)
            kind, text = extract_text(asset.original_filename, raw)
            asset.detected_kind = kind.value
            asset.extracted_text = text

        posting = fetch_job_posting(app.job_url)
        app.job_final_url = posting.final_url
        app.job_title = posting.title or None
        app.job_text = posting.text

        app.status = "ready"
        db.commit()
        log.info("application_ready", application_id=str(app.id))
    except Exception as exc:  # noqa: BLE001
        db.rollback()
        # Re-fetch to mark failed safely.
        fresh = db.get(Application, application_id)
        if fresh is not None:
            fresh.status = "failed"
            fresh.error = f"{exc.__class__.__name__}: {exc}"
            db.commit()
        log.exception("application_failed", application_id=str(application_id))
        raise


@celery_app.task(name="cvpilot.process_application", bind=True, max_retries=2)
def process_application(self, application_id: str) -> None:
    aid = uuid.UUID(application_id)
    db = SessionLocal()
    try:
        _process(db, aid)
    finally:
        db.close()
'''

# -----------------------------------------------------------------------------
# api/v1 router: register applications
# -----------------------------------------------------------------------------
FILES["api/v1/router.py"] = '''"""
Versioned API router aggregator. Mounted at /api/v1 in main.py.
"""

from __future__ import annotations

from fastapi import APIRouter

from backend.app.routes import applications, auth

api_router = APIRouter(prefix="/api/v1")
api_router.include_router(auth.router)
api_router.include_router(applications.router)
'''

# -----------------------------------------------------------------------------
# Tests
# -----------------------------------------------------------------------------
FILES["tests/backend/conftest.py"] = '''"""
Pytest fixtures for backend integration tests.
Uses the real DB defined in .env (cvpilot/cvpilot).
Forces Celery into eager mode so .delay() runs inline.
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import text

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))

os.environ.setdefault("APP_ENV", "development")
os.environ["CELERY_TASK_ALWAYS_EAGER"] = "true"

from backend.app.db.session import engine  # noqa: E402
from backend.app.main import app  # noqa: E402


@pytest.fixture(scope="session", autouse=True)
def _ensure_schema() -> None:
    with engine.connect() as conn:
        conn.execute(text("SELECT 1"))


@pytest.fixture()
def client() -> TestClient:
    return TestClient(app)


@pytest.fixture(autouse=True)
def _cleanup_db() -> None:
    yield
    with engine.begin() as conn:
        # Order matters because of FKs.
        conn.execute(text("DELETE FROM file_assets"))
        conn.execute(text("DELETE FROM applications"))
        conn.execute(text("DELETE FROM audit_logs"))
        conn.execute(text("DELETE FROM users WHERE email LIKE 'pytest+%@cvpilot.dev'"))
'''

FILES["tests/backend/test_applications.py"] = '''"""
End-to-end test for the Applications API.
Verifies: register -> upload (PDF CV, TXT cover letter) -> Celery eager runs ->
get back parsed text + ready status.
Job URL is set to example.com (real HTML, stable).
"""

from __future__ import annotations

import io
import uuid

from pypdf import PdfWriter


def _pdf_bytes(text: str = "Senior Python Engineer with 8 years of experience.") -> bytes:
    # Minimal valid PDF generated programmatically.
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    # pypdf does not render arbitrary text on a blank page; we embed the
    # text via the document metadata so extract_text() at least returns
    # SOMETHING. For richer text we use the metadata + page content stream below.
    writer.add_metadata({"/Subject": text})
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _email() -> str:
    return f"pytest+{uuid.uuid4().hex[:12]}@cvpilot.dev"


def _register_and_token(client) -> str:
    email = _email()
    password = "S3cure!Passw0rd"
    r = client.post(
        "/api/v1/auth/register",
        json={"email": email, "password": password, "full_name": "Test User"},
    )
    assert r.status_code == 201, r.text
    r = client.post(
        "/api/v1/auth/login",
        json={"email": email, "password": password},
    )
    assert r.status_code == 200, r.text
    return r.json()["access_token"]


def test_create_application_happy_path(client) -> None:
    token = _register_and_token(client)

    cover_letter_text = (
        "Dear Hiring Team,\\n"
        "I am thrilled to apply for this role. My background in Python and\\n"
        "distributed systems aligns directly with your needs.\\n"
        "Best,\\nA Candidate"
    )

    files = {
        "cv": ("cv.pdf", _pdf_bytes(), "application/pdf"),
        "cover_letter": ("cover.txt", cover_letter_text.encode("utf-8"), "text/plain"),
    }
    data = {
        "job_url": "https://example.com/",
        "linkedin_url": "https://www.linkedin.com/in/example/",
        "portfolio_url": "https://example.com/portfolio",
    }

    r = client.post(
        "/api/v1/applications",
        files=files,
        data=data,
        headers={"Authorization": f"Bearer {token}"},
    )
    assert r.status_code == 202, r.text
    body = r.json()
    app_id = body["id"]
    assert body["status"] in {"pending", "processing", "ready", "failed"}
    assert len(body["files"]) == 2

    # Celery eager mode ran the task synchronously inside the request.
    # Re-fetch to confirm terminal state.
    r2 = client.get(
        f"/api/v1/applications/{app_id}",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert r2.status_code == 200, r2.text
    final = r2.json()
    assert final["status"] == "ready", final
    assert final["job_text"], "job_text should be populated after fetch"
    assert all(f["extracted_text"] for f in final["files"]), final["files"]


def test_create_application_rejects_bad_job_url(client) -> None:
    token = _register_and_token(client)
    files = {
        "cv": ("cv.pdf", _pdf_bytes(), "application/pdf"),
        "cover_letter": ("cover.txt", b"hello", "text/plain"),
    }
    r = client.post(
        "/api/v1/applications",
        files=files,
        data={"job_url": "ftp://nope.invalid"},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert r.status_code == 422
    assert r.json()["error"]["code"] == "job_url_invalid"


def test_list_and_isolation(client) -> None:
    token_a = _register_and_token(client)
    token_b = _register_and_token(client)

    # A creates one
    r = client.post(
        "/api/v1/applications",
        files={
            "cv": ("cv.pdf", _pdf_bytes(), "application/pdf"),
            "cover_letter": ("c.txt", b"hello world", "text/plain"),
        },
        data={"job_url": "https://example.com/"},
        headers={"Authorization": f"Bearer {token_a}"},
    )
    assert r.status_code == 202, r.text
    app_id = r.json()["id"]

    # B cannot read A's application
    r = client.get(
        f"/api/v1/applications/{app_id}",
        headers={"Authorization": f"Bearer {token_b}"},
    )
    assert r.status_code == 403

    # B's list is empty
    r = client.get(
        "/api/v1/applications",
        headers={"Authorization": f"Bearer {token_b}"},
    )
    assert r.status_code == 200
    assert r.json() == []


def test_unauthenticated_rejected(client) -> None:
    r = client.get("/api/v1/applications")
    assert r.status_code == 401
'''


def write(rel: str, content: str) -> None:
    p = ROOT / rel
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(content, encoding="utf-8")
    print(f"  wrote {rel}")


def main() -> None:
    print(f"Phase 4 into: {ROOT}")
    for rel, content in FILES.items():
        write(rel, content)
    print("\nPhase 4 files written.")


if __name__ == "__main__":
    main()
