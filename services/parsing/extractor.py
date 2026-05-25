"""
File-content extractor.
- PDF: pypdf
- DOCX: python-docx
- TXT: utf-8 decode with errors="ignore"
Validates magic bytes so a renamed .exe cannot smuggle past the extension check.
"""

from __future__ import annotations

import io
from enum import Enum

from docx import Document  # python-docx
from pypdf import PdfReader

from backend.app.core.errors import ValidationAppError


class FileKind(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


_PDF_MAGIC = b"%PDF-"
_DOCX_MAGIC = b"PK\x03\x04"  # any zip; we rely on python-docx to validate structure


def detect_kind(filename: str, data: bytes) -> FileKind:
    name = filename.lower()
    head = data[:8]

    if head.startswith(_PDF_MAGIC):
        return FileKind.PDF
    if head.startswith(b"PK\x03\x04") and name.endswith(".docx"):
        return FileKind.DOCX
    if name.endswith(".txt"):
        # Best-effort: assume text if it decodes mostly as utf-8.
        try:
            data[:4096].decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ValidationAppError(
                "File does not appear to be valid UTF-8 text.",
                code="invalid_file_encoding",
            ) from exc
        return FileKind.TXT

    raise ValidationAppError(
        "Unsupported or malformed file. Allowed: PDF, DOCX, TXT.",
        code="unsupported_file_type",
    )


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            parts.append("")
    return "\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def _extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore").strip()


def extract_text(filename: str, data: bytes) -> tuple[FileKind, str]:
    kind = detect_kind(filename, data)
    if kind == FileKind.PDF:
        text = _extract_pdf(data)
    elif kind == FileKind.DOCX:
        text = _extract_docx(data)
    else:
        text = _extract_txt(data)

    if not text:
        raise ValidationAppError(
            "Could not extract any text from the uploaded file.",
            code="empty_extraction",
        )
    return kind, text
